import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def _sanitize_content(content):
    if not content: return ""
    preamble_pattern = re.compile(r"^([\s\*\-_>]*)(To perform|I will|Sure|I'll|Certainly|Here is|Then, I'll proceed|In order to|Okay|I've|I can|I've noticed|First|I will first|Secondly|Let me)[\s\S]+?(\.|:|\n)", re.IGNORECASE | re.MULTILINE)
    cleaned = content.strip()
    last = ""
    pass_count = 0
    while cleaned != last and pass_count < 20:
        pass_count += 1
        last = cleaned
        cleaned = preamble_pattern.sub('', cleaned).strip()
    return cleaned
    
def _setup_page_layout(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    pass

def _add_title_block(doc, title):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.0)
    
    c00 = table.cell(0, 0)
    p = c00.paragraphs[0]
    r = p.add_run("ELDORIA STRATEGIC ANALYSIS")
    r.font.name = 'Arial'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.font.all_caps = True
    p.paragraph_format.space_after = Pt(0)
    
    c01 = table.cell(0, 1)
    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TIMESTAMP")
    r.font.name = 'Arial'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.font.all_caps = True
    p.paragraph_format.space_after = Pt(0)
    
    c10 = table.cell(1, 0)
    p = c10.paragraphs[0]
    r = p.add_run(title)
    r.font.name = 'Arial'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)
    p.paragraph_format.space_before = Pt(4)
    
    c11 = table.cell(1, 1)
    p = c11.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(datetime.now().strftime("%B %d, %Y"))
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_before = Pt(12)

    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(12)
    
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '06B6D4')
    pbdr.append(bottom)
    pPr.append(pbdr)

    doc.add_paragraph()

def _add_footer(doc, page_num=True):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def add_run(text):
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        return r

    # helper to create field run
    def create_field_run(instr):
        r = p.add_run()
        r.font.name = 'Arial' 
        r.font.size = Pt(8)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instr
        r._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r._r.append(fldChar2)

    add_run("Page ")
    create_field_run("PAGE")
    add_run(" of ")
    create_field_run("NUMPAGES")
    add_run(" | GENERATED VIA ELDORIA AI IDE")

def _create_vertical_bar_header(doc, text, level=2):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    col0_w = Inches(0.05)
    col1_w = Inches(6.0)
    
    table.columns[0].width = col0_w
    table.columns[1].width = col1_w
    
    cell0 = table.cell(0, 0)
    cell0.width = col0_w
    _set_shading(cell0._tc.get_or_add_tcPr(), "007BFF")
    
    cell1 = table.cell(0, 1)
    cell1.width = col1_w
    p = cell1.paragraphs[0]
    run = p.add_run(text.upper())
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)
    run.font.size = Pt(14 if level > 1 else 16)
    
    doc.add_paragraph()

def _setup_styles(doc):
    """Configures document styles to match Eldoria's web aesthetic (Cyan/Sans-Serif)"""
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    for i in range(1, 4):
        style_name = f'Heading {i}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, 1)
            
        style.font.name = 'Arial'
        style.font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)
        if i == 1: style.font.size = Pt(18)
        if i == 2: style.font.size = Pt(16)
        if i == 3: style.font.size = Pt(14)

def _set_shading(element, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    element.append(shading_elm)

def _add_markdown_content(doc, content):
    """
    Parses basic Markdown (Headers, Lists, Tables, Bold) into Docx elements.
    Removes SAF blocks.
    """
    if not content: return
    content = _sanitize_content(content)
    
    # Transform SAF blocks (Match Print View)
    content = re.sub(r"```json\s*<SAF_ISO>", "```json", content, flags=re.IGNORECASE)
    content = re.sub(r"</SAF_ISO>\s*```", "```", content, flags=re.IGNORECASE)
    content = re.sub(r"<SAF_ISO>", "\n\n### Technical Specification (SAF-ISO)\n```json\n", content, flags=re.IGNORECASE)
    content = re.sub(r"</SAF_ISO>", "\n```\n", content, flags=re.IGNORECASE)

    lines = content.split('\n')
    in_code_block = False
    table_buffer = []

    def flush_table():
        if not table_buffer: return
        rows = len(table_buffer)
        cols = len(table_buffer[0].split('|')) - 2
        if cols < 1: return
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for r_idx, row_str in enumerate(table_buffer):
            cells = [c.strip() for c in row_str.split('|')[1:-1]]
            for c_idx, cell_text in enumerate(cells):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    if r_idx == 0:
                         _set_shading(cell._tc.get_or_add_tcPr(), "F8FAFC")
                         for paragraph in cell.paragraphs:
                             for run in paragraph.runs:
                                 run.font.bold = True
        doc.add_paragraph()
        table_buffer.clear()

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            continue
        elif table_buffer:
             flush_table()

        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            t = doc.add_table(rows=1, cols=1)
            cell = t.cell(0, 0)
            _set_shading(cell._tc.get_or_add_tcPr(), "F1F5F9")
            p = cell.paragraphs[0]
            p.text = stripped
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)
            continue

        if stripped.startswith('# '): _create_vertical_bar_header(doc, stripped[2:], level=1)
        elif stripped.startswith('## '): _create_vertical_bar_header(doc, stripped[3:], level=2)
        elif stripped.startswith('### '): _create_vertical_bar_header(doc, stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            try: p = doc.add_paragraph(style='List Bullet')
            except: p = doc.add_paragraph(style='List Paragraph')
            _process_bold(p, stripped[2:])
        else:
            if not stripped: continue
            p = doc.add_paragraph()
            _process_bold(p, stripped)
    
    if table_buffer: flush_table()

def _process_bold(paragraph, text):
    segments = text.split('**')
    for i, segment in enumerate(segments):
        if not segment: continue
        run = paragraph.add_run(segment)
        if i % 2 != 0: run.bold = True

def build_thesis(input_data):
    doc = Document()
    _setup_page_layout(doc)
    _setup_styles(doc)
    _add_header(doc, right_text="ACADEMIC DRAFT")
    _add_footer(doc)
    
    # 1. Title Page
    title = input_data.get('basics', {}).get('title', 'UNTITLED THESIS').upper()
    author = input_data.get('basics', {}).get('author', 'UNKNOWN AUTHOR')
    reg_num = input_data.get('basics', {}).get('regNumber', '')
    year = input_data.get('basics', {}).get('year', '2024')
    
    # Simple Title Page Logic
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(16)
    
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph("BY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(author)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].bold = True
    
    p = doc.add_paragraph(reg_num)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph("A THESIS SUBMITTED TO THE DEPARTMENT OF MECHANICAL ENGINEERING")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("RIVERS STATE UNIVERSITY, PORT HARCOURT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(year)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 2. Content Sections
    drafts = input_data.get('draft_content', {})
    for chapter_name, content in drafts.items():
        doc.add_heading(chapter_name, level=1)
        _add_markdown_content(doc, content)
        doc.add_page_break()
    
    output_path = f"thesis_{input_data.get('id', 'temp')}.docx"
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Error: No input JSON provided")
        sys.exit(1)
        
    try:
        data = json.loads(sys.argv[1])
        path = build_thesis(data)
        print(f"SUCCESS: {path}")
    except Exception as e:
        print(f"ERROR: {str(e)}")
        sys.exit(1)
