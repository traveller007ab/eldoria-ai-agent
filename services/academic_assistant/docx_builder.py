import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
from datetime import datetime
import re

def _sanitize_content(content):
    """
    Robustly strips AI preambles using a multi-pass regex loop.
    Matches the logic in OutputPanel.tsx and AcademicHub.tsx.
    """
    if not content:
        return ""
        
    # Aggressive preamble regex
    # Handles Markdown prefixes (e.g. **Here is...) and colons, AND newlines ([\s\S])
    # Note: In Python, re.DOTALL makes '.' match newlines, but we use [\s\S] equivalent logic or just standard multiline
    
    preamble_pattern = re.compile(r"^([\s\*\-_>]*)(To perform|I will|Sure|I'll|Certainly|Here is|Then, I'll proceed|In order to|Okay|I've|I can|I've noticed|First|I will first|Secondly|Let me)[\s\S]+?(\.|:|\n)", re.IGNORECASE | re.MULTILINE)
    
    cleaned_content = content.strip()
    last_cleaned = ""
    
    # Multi-pass loop to catch consecutive/nested preambles
    pass_count = 0
    while cleaned_content != last_cleaned and pass_count < 20:
        pass_count += 1
        last_cleaned = cleaned_content
        # re.sub replaces all occurrences, but since we are anchoring to ^ (start of line), 
        # it effectively removes leading preambles.
        cleaned_content = preamble_pattern.sub('', cleaned_content).strip()
        
    return cleaned_content

def _setup_page_layout(doc):
    """Applies standard A4/Letter margins (1 inch)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Remove default header content, we will build a custom Title Block in body
    # section = doc.sections[0]
    # header = section.header
    # for p in header.paragraphs: p.text = "" 
    # (Actually, keep the simple header if needed, but the image implies the Title Block IS the header or top of page)
    # Let's keep the timestamp/title header as a fallback for subsequent pages, 
    # but for Page 1, we build the "Big" block.
    pass

def _add_title_block(doc, title, project_id=""):
    """
    Creates a PREMIUM 'Strategic Analysis' title block.
    - Full-width cyan accent bar at top
    - Shaded title container
    - Professional typography hierarchy
    """
    # === 1. FULL-WIDTH ACCENT BAR (Letterhead Style) ===
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.autofit = False
    accent_table.allow_autofit = False
    accent_table.columns[0].width = Inches(6.5)
    accent_cell = accent_table.cell(0, 0)
    accent_cell.width = Inches(6.5)
    _set_shading(accent_cell._tc.get_or_add_tcPr(), "0891B2")  # Cyan-600
    
    # Add small text inside the bar
    p = accent_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("  ELDORIA AI IDE  •  STRATEGIC ANALYSIS FRAMEWORK  •  NEURAL CONTEXT LAYER v1.2")
    r.font.name = 'Arial'
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White
    r.font.bold = True
    
    # Small spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    
    # === 2. MAIN TITLE BLOCK WITH SHADED BACKGROUND ===
    title_table = doc.add_table(rows=2, cols=2)
    title_table.autofit = False
    title_table.allow_autofit = False
    
    # Widths
    title_table.columns[0].width = Inches(4.8)
    title_table.columns[1].width = Inches(1.7)
    
    # Apply subtle background to all cells
    for row in title_table.rows:
        for cell in row.cells:
            _set_shading(cell._tc.get_or_add_tcPr(), "F8FAFC")  # Slate-50
    
    # --- Row 1: Labels ---
    c00 = title_table.cell(0, 0)
    p = c00.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("STRATEGIC BRIEF")
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # Slate-500
    r.font.bold = True
    
    c01 = title_table.cell(0, 1)
    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("GENERATED")
    r.font.name = 'Arial'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # Slate-400
    
    # --- Row 2: Title & Date ---
    c10 = title_table.cell(1, 0)
    p = c10.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    r.font.name = 'Arial'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # Slate-900 (Dark, not blue)
    
    c11 = title_table.cell(1, 1)
    p = c11.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(datetime.now().strftime("%b %d, %Y"))
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    r.font.bold = True
    
    # === 3. THIN ACCENT LINE BELOW ===
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(16)
    
    pPr = line_para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')  # 2.25pt - thicker
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0891B2')  # Cyan-600
    pbdr.append(bottom)
    pPr.append(pbdr)

def _add_footer(doc, page_num=True):
    """Adds footer with Page X of Y and generation info"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # helper for gray text runs
    def add_run(text):
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        return r

    # helper to create field run
    def create_field_run(instr):
        r = p.add_run()
        r.font.name = 'Arial' 
        r.font.size = Pt(8)
        
        # Begin
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r._r.append(fldChar1)
        
        # Instruction
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instr
        r._r.append(instrText)
        
        # End
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r._r.append(fldChar2)

    add_run("Page ")
    create_field_run("PAGE")
    add_run(" of ")
    create_field_run("NUMPAGES")
    add_run(" | GENERATED VIA ELDORIA AI IDE")

def _create_vertical_bar_header(doc, text, level=2):
    """
    PREMIUM section header with colored accent bar and professional styling.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False 
    table.allow_autofit = False
    
    # Column widths
    col0_w = Inches(0.08)  # Slightly thicker bar
    col1_w = Inches(6.42)
    
    table.columns[0].width = col0_w
    table.columns[1].width = col1_w
    
    # Accent bar cell
    cell0 = table.cell(0, 0)
    cell0.width = col0_w
    _set_shading(cell0._tc.get_or_add_tcPr(), "0891B2")  # Cyan-600 (matches header bar)
    
    # Text cell with subtle background
    cell1 = table.cell(0, 1)
    cell1.width = col1_w
    _set_shading(cell1._tc.get_or_add_tcPr(), "F1F5F9")  # Slate-100
    
    p = cell1.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run("  " + text.upper())  # Small indent
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # Slate-900 (dark, professional)
    run.font.size = Pt(13 if level > 1 else 15)
    
    # Tighter spacing after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

def _setup_styles(doc):
    """Configures document styles to match Eldoria's web aesthetic (Blue/Sans-Serif)"""
    # 1. Normal Text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0) # Force Black to prevent style bleeding
    
    # 2. Headings (Blue #007BFF)
    for i in range(1, 4):
        style_name = f'Heading {i}'
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, 1) # 1 = Paragraph Style
            
        font = style.font
        font.name = 'Arial'
        font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)
        if i == 1: font.size = Pt(18) # Web H1
        if i == 2: font.size = Pt(16) # Web H2
        if i == 3: font.size = Pt(14) # Web H3
def _set_shading(element, color_hex):
    """Helper to add background color to a table cell or paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    element.append(shading_elm)

def _add_markdown_content(doc, content):
    """
    Parses basic Markdown (Headers, Lists, Tables, Bold) into Docx elements.
    Removes SAF blocks.
    """
    if not content: return

    # 1. Sanitize preambles first
    content = _sanitize_content(content)

    # 2. Transform SAF/JSON blocks to match Print View ("Technical Specification")
    # Instead of deleting, we reformat to a Header + Code Block
    
    # Handle wrapped cases first: ```json<SAF_ISO> -> ```json
    content = re.sub(r"```json\s*<SAF_ISO>", "```json", content, flags=re.IGNORECASE)
    content = re.sub(r"</SAF_ISO>\s*```", "```", content, flags=re.IGNORECASE)
    
    # Handle bare tags: <SAF_ISO> -> ### Header + ```json
    content = re.sub(r"<SAF_ISO>", "\n\n### Technical Specification (SAF-ISO)\n```json\n", content, flags=re.IGNORECASE)
    content = re.sub(r"</SAF_ISO>", "\n```\n", content, flags=re.IGNORECASE)

    lines = content.split('\n')
    
    # State flags
    in_code_block = False
    table_buffer = []
    
    def flush_table():
        if not table_buffer: return
        # Create table
        rows = len(table_buffer)
        cols = len(table_buffer[0].split('|')) - 2 # Assuming | val | val | format
        if cols < 1: return
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for r_idx, row_str in enumerate(table_buffer):
            # Split and clean empty start/end
            cells = [c.strip() for c in row_str.split('|')[1:-1]]
            for c_idx, cell_text in enumerate(cells):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    if r_idx == 0: # Header shading
                         _set_shading(cell._tc.get_or_add_tcPr(), "F8FAFC")
                         for paragraph in cell.paragraphs:
                             for run in paragraph.runs:
                                 run.font.bold = True
        
        doc.add_paragraph() # Spacer
        table_buffer.clear()

    for line in lines:
        stripped = line.strip()
        
        # Table Detection
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            continue
        elif table_buffer:
             flush_table()

        # Code block toggle
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            # Render code inside a shaded table cell for background color
            t = doc.add_table(rows=1, cols=1)
            cell = t.cell(0, 0)
            _set_shading(cell._tc.get_or_add_tcPr(), "F1F5F9") # Slate-100
            p = cell.paragraphs[0]
            p.text = stripped
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)
            continue

        # Headers - Use Special Vertical Bar Style
        if stripped.startswith('# '):
            _create_vertical_bar_header(doc, stripped[2:], level=1)
        elif stripped.startswith('## '):
            _create_vertical_bar_header(doc, stripped[3:], level=2)
        elif stripped.startswith('### '):
            _create_vertical_bar_header(doc, stripped[4:], level=3)
        
        # List Items - Bullet
        elif stripped.startswith('- ') or stripped.startswith('* '):
            try:
                p = doc.add_paragraph(style='List Bullet')
            except:
                p = doc.add_paragraph(style='List Paragraph') # Fallback
            _process_bold(p, stripped[2:])
            
        # List Items - Numbered
        elif re.match(r'^\d+\.\s', stripped):
            try:
                p = doc.add_paragraph(style='List Number')
            except:
                 p = doc.add_paragraph(style='List Paragraph')
            # Remove "1. " prefix
            text_content = re.sub(r'^\d+\.\s', '', stripped)
            _process_bold(p, text_content)
            
        # Standard Paragraph
        else:
            if not stripped: continue
            p = doc.add_paragraph()
            _process_bold(p, stripped)
    
    # Flush any remaining table
    if table_buffer: flush_table()


def _process_bold(paragraph, text):
    """Refactored bold processing"""
    segments = text.split('**')
    for i, segment in enumerate(segments):
        if not segment: continue
        run = paragraph.add_run(segment)
        if i % 2 != 0: # Odd indices are inside **
            run.bold = True

def build_thesis(input_data):
    doc = Document()
    _setup_page_layout(doc)
    _setup_styles(doc) # Keep valid for fallbacks
    # For thesis we don't need a running header, just the footer
    _add_footer(doc)
    
    # 1. Title Page
    title = input_data.get('basics', {}).get('title', 'UNTITLED THESIS').upper()
    author = input_data.get('basics', {}).get('author', 'UNKNOWN AUTHOR')
    reg_num = input_data.get('basics', {}).get('regNumber', '')
    year = input_data.get('basics', {}).get('year', '2024')
    
    # Simple Title Page Logic
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph("BY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(author)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph(reg_num)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph("A THESIS SUBMITTED TO THE DEPARTMENT OF MECHANICAL ENGINEERING")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("RIVERS STATE UNIVERSITY, PORT HARCOURT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(year)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 2. Content Sections
    ordered_chapters = [
        'Front Matter',
        'Abstract', 
        'Chapter 1: Introduction', 
        'Chapter 2: Literature Review', 
        'Chapter 3: Materials & Methods',
        'Chapter 4: Results & Discussion',
        'Chapter 5: Conclusion & Recommendations'
    ]
    
    drafts = input_data.get('draft_content', {})
    for chapter_name in ordered_chapters:
        content = drafts.get(chapter_name)
        if content:
            # We can use the simple header or just a page break + heading
            # For consistency with the Thesis look, let's just use the markdown parser
            # The parser handles H1/H2 etc.
            # If we want a specific "Title Page" for each chapter, we could do:
            # _add_title_block(doc, chapter_name) 
            # But that might be too heavy. Let's stick to standard flow:
            
            # _add_markdown_content will parse the headers in the content.
            # If the content doesn't have a header, we might want to add one.
            # But safely, let's just dump the content.
            
            _add_markdown_content(doc, content)
            doc.add_page_break()

    # 3. References
    references = input_data.get('references', [])
    if references:
        doc.add_heading('REFERENCES', level=1)
        for ref in references:
            apa = ref.get('formattedApa', f"{ref.get('authors')} ({ref.get('year')}). {ref.get('title')}. {ref.get('journal')}.")
            p = doc.add_paragraph(apa)
            p.style.font.size = Pt(11)
    
    # 4. Ethical Watermark (Footer)
    # Footer handled by _add_footer

    output_path = f"thesis_{input_data.get('id', 'temp')}.docx"
    doc.save(output_path)
    return output_path

def build_simple_doc(title, content):
    doc = Document()
    _setup_page_layout(doc)
    _setup_styles(doc) # Fallbacks
    _add_title_block(doc, title)
    _add_footer(doc)
    
    # Body Content via Parser
    _add_markdown_content(doc, content)
                
    # Footer is handled by _add_footer

    filename = f"export_{int(time.time())}.docx"
    output_path = os.path.join(os.getcwd(), filename)
    doc.save(output_path)
    return filename, output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Error: No input JSON provided")
        sys.exit(1)
        
    try:
        data = json.loads(sys.argv[1])
        path = build_thesis(data)
        print(f"SUCCESS: {path}")
    except Exception as e:
        print(f"ERROR: {str(e)}")
        sys.exit(1)
